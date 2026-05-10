# import re
# from pathlib import Path
# from typing import Dict, List, Optional

# import fitz
# from groq import Groq
# from langchain_core.documents import Document

# from src.chunker import chunk_documents
# from src.vector_store import (
#     build_vector_store_from_chunks,
#     get_vector_store,
#     reset_vector_store,
# )
# from src.session_manager import create_user_folders
# from src.config import (
#     get_groq_api_key,
#     get_groq_model,
#     get_groq_temperature,
#     get_groq_max_tokens,
#     is_groq_configured,
#     CHUNK_SIZE,
#     CHUNK_OVERLAP,
# )


# RESUME_SYSTEM_PROMPT = """
# You are an ATS Resume and Job Description RAG Analyzer.

# You compare a resume with a job description using retrieved resume/JD context.

# Rules:
# 1. Use only the given resume text, JD text, and retrieved RAG context.
# 2. Do not assume missing skills.
# 3. Do not fake projects, experience, tools, or achievements.
# 4. Give honest match percentage.
# 5. Clearly list matching skills.
# 6. Clearly list missing skills.
# 7. Tell whether the user should apply or not.
# 8. Suggest better alternative roles if this JD is not a strong match.
# 9. Suggest resume improvements based on the JD.
# 10. Use simple English.
# 11. Make the output useful for students and freshers.
# 12. Keep the answer concise and useful.
# """


# RESUME_ANALYSIS_PROMPT = """
# Resume Text:
# {resume_text}

# Job Description Text:
# {jd_text}

# Retrieved RAG Context:
# {rag_context}

# Analyze resume against the job description.

# Important:
# - Use resume text to check what the user has.
# - Use JD text to check what the company needs.
# - Use retrieved RAG context to support matching.
# - Do not add fake skills.
# - Do not say the user has a skill if it is not in the resume.
# - Keep the answer practical and not too long.

# Give answer in this exact format:

# Match Percentage:
# - Give one percentage from 0 to 100.
# - Explain the reason shortly.

# Verdict:
# - Strong Match / Good Match / Average Match / Weak Match / Not Recommended

# Should You Apply:
# - Yes / Maybe / No
# - Give short reason.

# Matching Skills:
# - List skills/projects/experience from resume that match the JD.

# Missing Skills:
# - List important JD skills missing from resume.

# Weak Areas:
# - List weak points in resume for this JD.

# Resume Improvement Suggestions:
# - Give practical improvement points.

# Best Alternative Roles:
# - Suggest better job roles based on resume.

# Alternative Resume Direction:
# - Suggest how the resume should be positioned.

# ATS-Friendly Summary:
# - Write 3 to 4 lines for this JD.
# - Do not add fake skills.

# Final Advice:
# - Give short final advice.
# """


# def clean_text(text: str) -> str:
#     if not text:
#         return ""

#     text = text.replace("\x00", " ")
#     text = text.replace("\t", " ")
#     text = text.replace("\r", "\n")

#     lines = []

#     for line in text.splitlines():
#         line = line.strip()
#         if line:
#             lines.append(line)

#     return "\n".join(lines).strip()


# def limit_text(text: str, max_chars: int = 5000) -> str:
#     text = clean_text(text)

#     if len(text) <= max_chars:
#         return text

#     return text[:max_chars]


# def extract_text_from_pdf_bytes(file_bytes: bytes) -> str:
#     if not file_bytes:
#         return ""

#     text_parts = []

#     with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
#         for page_index in range(pdf.page_count):
#             page = pdf.load_page(page_index)
#             text = clean_text(page.get_text("text"))

#             if text:
#                 text_parts.append(text)

#     return clean_text("\n\n".join(text_parts))


# def extract_text_from_txt_bytes(file_bytes: bytes) -> str:
#     if not file_bytes:
#         return ""

#     try:
#         return clean_text(file_bytes.decode("utf-8"))
#     except Exception:
#         return clean_text(file_bytes.decode("latin-1", errors="ignore"))


# def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
#     try:
#         import docx
#     except Exception:
#         raise ImportError("python-docx is not installed. Run: pip install python-docx")

#     import tempfile

#     with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
#         temp_file.write(file_bytes)
#         temp_path = temp_file.name

#     try:
#         doc = docx.Document(temp_path)
#         paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
#         return clean_text("\n".join(paragraphs))
#     finally:
#         try:
#             Path(temp_path).unlink(missing_ok=True)
#         except Exception:
#             pass


# def extract_text_from_uploaded_file(uploaded_file) -> Dict:
#     if uploaded_file is None:
#         return {
#             "success": False,
#             "text": "",
#             "filename": "",
#             "error": "No file uploaded.",
#         }

#     filename = getattr(uploaded_file, "name", "uploaded_file")
#     suffix = Path(filename).suffix.lower()

#     try:
#         file_bytes = bytes(uploaded_file.getbuffer())

#         if suffix == ".pdf":
#             text = extract_text_from_pdf_bytes(file_bytes)
#         elif suffix == ".txt":
#             text = extract_text_from_txt_bytes(file_bytes)
#         elif suffix == ".docx":
#             text = extract_text_from_docx_bytes(file_bytes)
#         else:
#             return {
#                 "success": False,
#                 "text": "",
#                 "filename": filename,
#                 "error": "Only PDF, DOCX, and TXT files are supported.",
#             }

#         text = clean_text(text)

#         if not text:
#             return {
#                 "success": False,
#                 "text": "",
#                 "filename": filename,
#                 "error": "Could not extract readable text from file.",
#             }

#         return {
#             "success": True,
#             "text": text,
#             "filename": filename,
#             "error": None,
#         }

#     except Exception as e:
#         return {
#             "success": False,
#             "text": "",
#             "filename": filename,
#             "error": str(e),
#         }


# def create_resume_rag_documents(
#     resume_text: str,
#     jd_text: str,
#     resume_filename: str = "resume",
#     jd_filename: str = "job_description",
# ) -> List[Document]:
#     documents = [
#         Document(
#             page_content=limit_text(resume_text, 10000),
#             metadata={
#                 "source_type": "resume",
#                 "pdf_name": resume_filename,
#                 "source": resume_filename,
#                 "page_number": 1,
#                 "page": 1,
#                 "content_type": "resume_text",
#             },
#         ),
#         Document(
#             page_content=limit_text(jd_text, 10000),
#             metadata={
#                 "source_type": "job_description",
#                 "pdf_name": jd_filename,
#                 "source": jd_filename,
#                 "page_number": 1,
#                 "page": 1,
#                 "content_type": "jd_text",
#             },
#         ),
#     ]

#     return documents


# def build_resume_rag_index(
#     base_session_id: str,
#     resume_text: str,
#     jd_text: str,
#     resume_filename: str = "resume",
#     jd_filename: str = "job_description",
# ) -> Dict:
#     resume_session_id = f"{base_session_id}_resume_match"

#     create_user_folders(resume_session_id)
#     reset_vector_store(resume_session_id)

#     documents = create_resume_rag_documents(
#         resume_text=resume_text,
#         jd_text=jd_text,
#         resume_filename=resume_filename,
#         jd_filename=jd_filename,
#     )

#     chunk_result = chunk_documents(
#         documents=documents,
#         session_id=resume_session_id,
#         chunk_size=CHUNK_SIZE,
#         chunk_overlap=CHUNK_OVERLAP,
#     )

#     if not chunk_result.get("success"):
#         return {
#             "success": False,
#             "session_id": resume_session_id,
#             "chunks": [],
#             "message": chunk_result.get("error", "Chunking failed."),
#         }

#     chunks = chunk_result.get("chunks", [])

#     store_result = build_vector_store_from_chunks(
#         session_id=resume_session_id,
#         chunks=chunks,
#         reset_before_add=True,
#     )

#     if not store_result.get("success"):
#         return {
#             "success": False,
#             "session_id": resume_session_id,
#             "chunks": chunks,
#             "message": store_result.get("message", "Vector store failed."),
#         }

#     return {
#         "success": True,
#         "session_id": resume_session_id,
#         "chunks": chunks,
#         "total_chunks": len(chunks),
#         "message": "Resume and JD indexed successfully.",
#     }


# def retrieve_resume_rag_context(
#     resume_session_id: str,
#     jd_text: str,
#     top_k: int = 4,
# ) -> Dict:
#     try:
#         vector_store = get_vector_store(resume_session_id)

#         query = (
#             "Resume JD match skills experience projects education tools technologies "
#             "missing skills weak areas role fit"
#         )

#         docs = vector_store.similarity_search(
#             query=query,
#             k=top_k,
#         )

#         context_parts = []

#         for i, doc in enumerate(docs, start=1):
#             metadata = doc.metadata or {}

#             source_type = metadata.get("source_type", "unknown")
#             source = metadata.get("source", "unknown")

#             context_parts.append(
#                 f"""
# [Context {i}]
# Type: {source_type}
# Source: {source}

# {limit_text(doc.page_content, 1200)}
# """
#             )

#         return {
#             "success": True,
#             "docs": docs,
#             "context": "\n\n".join(context_parts).strip(),
#             "retrieved_chunks": len(docs),
#         }

#     except Exception as e:
#         return {
#             "success": False,
#             "docs": [],
#             "context": "",
#             "retrieved_chunks": 0,
#             "error": str(e),
#         }


# def get_safe_max_tokens() -> int:
#     try:
#         return min(int(get_groq_max_tokens()), 900)
#     except Exception:
#         return 900


# def call_groq_resume_llm(
#     resume_text: str,
#     jd_text: str,
#     rag_context: str,
# ) -> str:
#     if not is_groq_configured():
#         raise ValueError("GROQ_API_KEY is missing. Add GROQ_API_KEY in your .env file.")

#     client = Groq(api_key=get_groq_api_key())

#     user_prompt = RESUME_ANALYSIS_PROMPT.format(
#         resume_text=limit_text(resume_text, 3000),
#         jd_text=limit_text(jd_text, 3000),
#         rag_context=limit_text(rag_context, 2500),
#     )

#     response = client.chat.completions.create(
#         model=get_groq_model(),
#         messages=[
#             {
#                 "role": "system",
#                 "content": RESUME_SYSTEM_PROMPT,
#             },
#             {
#                 "role": "user",
#                 "content": user_prompt,
#             },
#         ],
#         temperature=get_groq_temperature(),
#         max_tokens=get_safe_max_tokens(),
#     )

#     return response.choices[0].message.content.strip()


# def parse_match_data(answer: str) -> Dict:
#     data = {
#         "match_percentage": 0,
#         "verdict": "",
#         "should_apply": "",
#     }

#     if not answer:
#         return data

#     percent_match = re.search(r"(\d{1,3})\s*%", answer)

#     if percent_match:
#         percentage = int(percent_match.group(1))
#         data["match_percentage"] = max(0, min(100, percentage))

#     verdict_match = re.search(
#         r"Verdict\s*:\s*[-\s]*(Strong Match|Good Match|Average Match|Weak Match|Not Recommended)",
#         answer,
#         re.IGNORECASE,
#     )

#     if verdict_match:
#         data["verdict"] = verdict_match.group(1)

#     apply_match = re.search(
#         r"Should You Apply\s*:\s*[-\s]*(Yes|Maybe|No)",
#         answer,
#         re.IGNORECASE,
#     )

#     if apply_match:
#         data["should_apply"] = apply_match.group(1)

#     return data


# def analyze_resume_against_jd_with_rag(
#     base_session_id: str,
#     resume_text: str,
#     jd_text: str,
#     resume_filename: str = "resume",
#     jd_filename: str = "job_description",
# ) -> Dict:
#     resume_text = clean_text(resume_text)
#     jd_text = clean_text(jd_text)

#     if not resume_text:
#         return {
#             "success": False,
#             "answer": "Resume text is empty.",
#             "data": None,
#         }

#     if not jd_text:
#         return {
#             "success": False,
#             "answer": "Job description text is empty.",
#             "data": None,
#         }

#     try:
#         index_result = build_resume_rag_index(
#             base_session_id=base_session_id,
#             resume_text=resume_text,
#             jd_text=jd_text,
#             resume_filename=resume_filename,
#             jd_filename=jd_filename,
#         )

#         if not index_result.get("success"):
#             return {
#                 "success": False,
#                 "answer": index_result.get("message", "Resume RAG indexing failed."),
#                 "data": None,
#                 "index_result": index_result,
#             }

#         resume_session_id = index_result["session_id"]

#         retrieve_result = retrieve_resume_rag_context(
#             resume_session_id=resume_session_id,
#             jd_text=jd_text,
#             top_k=4,
#         )

#         if not retrieve_result.get("success"):
#             return {
#                 "success": False,
#                 "answer": retrieve_result.get("error", "Resume RAG retrieval failed."),
#                 "data": None,
#                 "index_result": index_result,
#                 "retrieve_result": retrieve_result,
#             }

#         rag_context = retrieve_result.get("context", "")

#         normal_answer = call_groq_resume_llm(
#             resume_text=resume_text,
#             jd_text=jd_text,
#             rag_context=rag_context,
#         )

#         parsed = parse_match_data(normal_answer)

#         return {
#             "success": True,
#             "answer": normal_answer,
#             "data": parsed,
#             "raw_json": "",
#             "resume_text": resume_text,
#             "jd_text": jd_text,
#             "rag_context": rag_context,
#             "resume_rag_session_id": resume_session_id,
#             "retrieved_chunks": retrieve_result.get("retrieved_chunks", 0),
#             "total_chunks": index_result.get("total_chunks", 0),
#             "model": get_groq_model(),
#             "provider": "groq",
#         }

#     except Exception as e:
#         return {
#             "success": False,
#             "answer": f"Error while analyzing resume with RAG: {str(e)}",
#             "data": None,
#             "model": get_groq_model(),
#             "provider": "groq",
#         }


# def analyze_uploaded_resume_and_jd_with_rag(
#     base_session_id: str,
#     resume_file,
#     jd_file=None,
#     jd_text: str = "",
# ) -> Dict:
#     resume_result = extract_text_from_uploaded_file(resume_file)

#     if not resume_result.get("success"):
#         return {
#             "success": False,
#             "answer": f"Resume error: {resume_result.get('error')}",
#             "data": None,
#             "resume_text": "",
#             "jd_text": "",
#         }

#     final_jd_text = ""
#     jd_filename = "pasted_job_description"

#     if jd_file is not None:
#         jd_result = extract_text_from_uploaded_file(jd_file)

#         if not jd_result.get("success"):
#             return {
#                 "success": False,
#                 "answer": f"JD error: {jd_result.get('error')}",
#                 "data": None,
#                 "resume_text": resume_result.get("text", ""),
#                 "jd_text": "",
#             }

#         final_jd_text = jd_result.get("text", "")
#         jd_filename = jd_result.get("filename", "job_description")
#     else:
#         final_jd_text = jd_text

#     analysis = analyze_resume_against_jd_with_rag(
#         base_session_id=base_session_id,
#         resume_text=resume_result.get("text", ""),
#         jd_text=final_jd_text,
#         resume_filename=resume_result.get("filename", "resume"),
#         jd_filename=jd_filename,
#     )

#     return analysis


# def render_resume_match_ui(st, session_id: str):
#     st.subheader("4️⃣ Resume vs JD Matcher Using RAG")

#     st.caption(
#         "Upload resume and job description. This uses extraction, chunking, ChromaDB retrieval, and Groq LLM analysis."
#     )

#     col1, col2 = st.columns(2)

#     with col1:
#         resume_file = st.file_uploader(
#             "Upload Resume",
#             type=["pdf", "docx", "txt"],
#             key="resume_match_resume_file",
#         )

#     with col2:
#         jd_file = st.file_uploader(
#             "Upload JD File Optional",
#             type=["pdf", "docx", "txt"],
#             key="resume_match_jd_file",
#         )

#     jd_text = st.text_area(
#         "Or paste Job Description here",
#         height=220,
#         placeholder="Paste job description here if you do not upload JD file...",
#         key="resume_match_jd_text",
#     )

#     analyze_clicked = st.button(
#         "Analyze Resume Match With RAG",
#         use_container_width=True,
#         key="resume_match_analyze_button",
#     )

#     if not analyze_clicked:
#         return

#     if resume_file is None:
#         st.error("Please upload your resume.")
#         return

#     if jd_file is None and not jd_text.strip():
#         st.error("Please upload JD file or paste JD text.")
#         return

#     with st.spinner("Extracting, indexing, retrieving, and analyzing with Groq..."):
#         result = analyze_uploaded_resume_and_jd_with_rag(
#             base_session_id=session_id,
#             resume_file=resume_file,
#             jd_file=jd_file,
#             jd_text=jd_text,
#         )

#     if not result.get("success"):
#         st.error(result.get("answer", "Analysis failed."))
#         return

#     data = result.get("data") or {}

#     percentage = data.get("match_percentage", 0)
#     verdict = data.get("verdict", "")
#     should_apply = data.get("should_apply", "")

#     col_a, col_b, col_c, col_d = st.columns(4)

#     with col_a:
#         st.metric("Match", f"{percentage}%")

#     with col_b:
#         st.metric("Verdict", verdict if verdict else "N/A")

#     with col_c:
#         st.metric("Apply?", should_apply if should_apply else "N/A")

#     with col_d:
#         st.metric("Provider", result.get("provider", "groq"))

#     st.caption(
#         f"RAG Session: {result.get('resume_rag_session_id')} | "
#         f"Chunks: {result.get('total_chunks')} | "
#         f"Retrieved: {result.get('retrieved_chunks')} | "
#         f"Model: {result.get('model')}"
#     )

#     st.markdown(result.get("answer", ""))

#     with st.expander("Retrieved RAG Context"):
#         st.text(result.get("rag_context", ""))

#     with st.expander("Extracted Resume Text"):
#         st.text(result.get("resume_text", ""))

#     with st.expander("Extracted JD Text"):
#         st.text(result.get("jd_text", ""))

#     with st.expander("Structured Result"):
#         st.json(data)


import re
from pathlib import Path
from typing import Dict, List, Optional

import fitz
from groq import Groq
from langchain_core.documents import Document

from src.chunker import chunk_documents
from src.vector_store import (
    build_vector_store_from_chunks,
    get_vector_store,
    reset_vector_store,
)

from src.session_manager import create_user_folders

from src.config import (
    get_groq_api_key,
    get_groq_model,
    get_groq_temperature,
    get_groq_max_tokens,
    is_groq_configured,
    CHUNK_SIZE,
    CHUNK_OVERLAP,
)

from src.prompts import (
    RESUME_GAP_SYSTEM_PROMPT,
    RESUME_GAP_PROMPT,
)


# ---------------------------------------------------------
# Resume Gap Finder
# Used by FastAPI route:
# POST /api/resume/gap-analysis
# ---------------------------------------------------------


RESUME_SYSTEM_PROMPT = RESUME_GAP_SYSTEM_PROMPT


RESUME_ANALYSIS_PROMPT = """
Resume Text:
{resume_text}

Job Description Text:
{jd_text}

Retrieved RAG Context:
{rag_context}

Analyze resume against the job description.

Important:
- Use resume text to check what the user has.
- Use JD text to check what the company needs.
- Use retrieved RAG context to support matching.
- Do not add fake skills.
- Do not say the user has a skill if it is not in the resume.
- Keep the answer practical and not too long.

Give answer in this exact format:

Match Percentage:
- Give one percentage from 0 to 100.
- Explain the reason shortly.

Verdict:
- Strong Match / Good Match / Average Match / Weak Match / Not Recommended

Should You Apply:
- Yes / Maybe / No
- Give short reason.

Matching Skills:
- List skills/projects/experience from resume that match the JD.

Missing Skills:
- List important JD skills missing from resume.

Weak Areas:
- List weak points in resume for this JD.

Resume Improvement Suggestions:
- Give practical improvement points.

Best Alternative Roles:
- Suggest better job roles based on resume.

Alternative Resume Direction:
- Suggest how the resume should be positioned.

ATS-Friendly Summary:
- Write 3 to 4 lines for this JD.
- Do not add fake skills.

Final Advice:
- Give short final advice.
""".strip()


# ---------------------------------------------------------
# Text cleaning helpers
# ---------------------------------------------------------

def clean_text(text: str) -> str:
    if not text:
        return ""

    text = str(text)

    text = text.replace("\x00", " ")
    text = text.replace("\t", " ")
    text = text.replace("\r", "\n")

    lines = []

    for line in text.splitlines():
        line = line.strip()

        if line:
            lines.append(line)

    return "\n".join(lines).strip()


def limit_text(text: str, max_chars: int = 5000) -> str:
    text = clean_text(text)

    if len(text) <= max_chars:
        return text

    return text[:max_chars]


# ---------------------------------------------------------
# File text extraction
# ---------------------------------------------------------

def extract_text_from_pdf_bytes(file_bytes: bytes) -> str:
    if not file_bytes:
        return ""

    text_parts = []

    with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
        for page_index in range(pdf.page_count):
            page = pdf.load_page(page_index)
            text = clean_text(page.get_text("text"))

            if text:
                text_parts.append(text)

    return clean_text("\n\n".join(text_parts))


def extract_text_from_txt_bytes(file_bytes: bytes) -> str:
    if not file_bytes:
        return ""

    try:
        return clean_text(file_bytes.decode("utf-8"))
    except Exception:
        return clean_text(file_bytes.decode("latin-1", errors="ignore"))


def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    try:
        import docx
    except Exception:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")

    import tempfile

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        temp_file.write(file_bytes)
        temp_path = temp_file.name

    try:
        doc = docx.Document(temp_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        return clean_text("\n".join(paragraphs))
    finally:
        try:
            Path(temp_path).unlink(missing_ok=True)
        except Exception:
            pass


def extract_text_from_file_bytes(
    file_bytes: bytes,
    filename: str = "uploaded_file",
) -> Dict:
    """
    FastAPI-friendly extractor.

    Supports:
    - PDF
    - TXT
    - DOCX
    """

    if not file_bytes:
        return {
            "success": False,
            "text": "",
            "filename": filename,
            "error": "File is empty.",
        }

    filename = filename or "uploaded_file"
    suffix = Path(filename).suffix.lower()

    try:
        if suffix == ".pdf":
            text = extract_text_from_pdf_bytes(file_bytes)
        elif suffix == ".txt":
            text = extract_text_from_txt_bytes(file_bytes)
        elif suffix == ".docx":
            text = extract_text_from_docx_bytes(file_bytes)
        else:
            return {
                "success": False,
                "text": "",
                "filename": filename,
                "error": "Only PDF, DOCX, and TXT files are supported.",
            }

        text = clean_text(text)

        if not text:
            return {
                "success": False,
                "text": "",
                "filename": filename,
                "error": "Could not extract readable text from file.",
            }

        return {
            "success": True,
            "text": text,
            "filename": filename,
            "error": None,
        }

    except Exception as e:
        return {
            "success": False,
            "text": "",
            "filename": filename,
            "error": str(e),
        }


def extract_text_from_uploaded_file(uploaded_file) -> Dict:
    """
    Old Streamlit-compatible extractor kept for safety.

    Streamlit uploaded_file has:
    - .name
    - .getbuffer()
    """

    if uploaded_file is None:
        return {
            "success": False,
            "text": "",
            "filename": "",
            "error": "No file uploaded.",
        }

    filename = getattr(uploaded_file, "name", "uploaded_file")

    try:
        file_bytes = bytes(uploaded_file.getbuffer())

        return extract_text_from_file_bytes(
            file_bytes=file_bytes,
            filename=filename,
        )

    except Exception as e:
        return {
            "success": False,
            "text": "",
            "filename": filename,
            "error": str(e),
        }


# ---------------------------------------------------------
# Resume + JD RAG documents
# ---------------------------------------------------------

def create_resume_rag_documents(
    resume_text: str,
    jd_text: str,
    resume_filename: str = "resume",
    jd_filename: str = "job_description",
) -> List[Document]:
    documents = [
        Document(
            page_content=limit_text(resume_text, 10000),
            metadata={
                "source_type": "resume",
                "pdf_name": resume_filename,
                "source": resume_filename,
                "page_number": 1,
                "page": 1,
                "content_type": "resume_text",
            },
        ),
        Document(
            page_content=limit_text(jd_text, 10000),
            metadata={
                "source_type": "job_description",
                "pdf_name": jd_filename,
                "source": jd_filename,
                "page_number": 1,
                "page": 1,
                "content_type": "jd_text",
            },
        ),
    ]

    return documents


def build_resume_rag_index(
    base_session_id: str,
    resume_text: str,
    jd_text: str,
    resume_filename: str = "resume",
    jd_filename: str = "job_description",
) -> Dict:
    resume_session_id = f"{base_session_id}_resume_match"

    create_user_folders(resume_session_id)
    reset_vector_store(resume_session_id)

    documents = create_resume_rag_documents(
        resume_text=resume_text,
        jd_text=jd_text,
        resume_filename=resume_filename,
        jd_filename=jd_filename,
    )

    chunk_result = chunk_documents(
        documents=documents,
        session_id=resume_session_id,
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
    )

    if not chunk_result.get("success"):
        return {
            "success": False,
            "session_id": resume_session_id,
            "chunks": [],
            "message": chunk_result.get("error", "Chunking failed."),
        }

    chunks = chunk_result.get("chunks", [])

    store_result = build_vector_store_from_chunks(
        session_id=resume_session_id,
        chunks=chunks,
        reset_before_add=True,
    )

    if not store_result.get("success"):
        return {
            "success": False,
            "session_id": resume_session_id,
            "chunks": chunks,
            "message": store_result.get("message", "Vector store failed."),
        }

    return {
        "success": True,
        "session_id": resume_session_id,
        "chunks": chunks,
        "total_chunks": len(chunks),
        "message": "Resume and JD indexed successfully.",
    }


# ---------------------------------------------------------
# Resume RAG retrieval
# ---------------------------------------------------------

def retrieve_resume_rag_context(
    resume_session_id: str,
    jd_text: str,
    top_k: int = 4,
) -> Dict:
    try:
        vector_store = get_vector_store(resume_session_id)

        query = (
            "Resume JD match skills experience projects education tools technologies "
            "missing skills weak areas role fit"
        )

        if jd_text:
            query = query + "\n" + limit_text(jd_text, 1000)

        docs = vector_store.similarity_search(
            query=query,
            k=top_k,
        )

        context_parts = []

        for i, doc in enumerate(docs, start=1):
            metadata = doc.metadata or {}

            source_type = metadata.get("source_type", "unknown")
            source = metadata.get("source", "unknown")

            context_parts.append(
                f"""
[Context {i}]
Type: {source_type}
Source: {source}

{limit_text(doc.page_content, 1200)}
""".strip()
            )

        return {
            "success": True,
            "docs": docs,
            "context": "\n\n".join(context_parts).strip(),
            "retrieved_chunks": len(docs),
        }

    except Exception as e:
        return {
            "success": False,
            "docs": [],
            "context": "",
            "retrieved_chunks": 0,
            "error": str(e),
        }


# ---------------------------------------------------------
# Groq LLM
# ---------------------------------------------------------

def get_safe_max_tokens() -> int:
    try:
        return min(int(get_groq_max_tokens()), 1200)
    except Exception:
        return 900


def call_groq_resume_llm(
    resume_text: str,
    jd_text: str,
    rag_context: str,
) -> str:
    if not is_groq_configured():
        raise ValueError(
            "GROQ_API_KEY is missing. Add GROQ_API_KEY in your root .env.local file."
        )

    client = Groq(api_key=get_groq_api_key())

    user_prompt = RESUME_ANALYSIS_PROMPT.format(
        resume_text=limit_text(resume_text, 3500),
        jd_text=limit_text(jd_text, 3500),
        rag_context=limit_text(rag_context, 2500),
    )

    response = client.chat.completions.create(
        model=get_groq_model(),
        messages=[
            {
                "role": "system",
                "content": RESUME_SYSTEM_PROMPT,
            },
            {
                "role": "user",
                "content": user_prompt,
            },
        ],
        temperature=get_groq_temperature(),
        max_tokens=get_safe_max_tokens(),
    )

    return response.choices[0].message.content.strip()


# ---------------------------------------------------------
# Parse useful structured values
# ---------------------------------------------------------

def parse_match_data(answer: str) -> Dict:
    data = {
        "match_percentage": 0,
        "verdict": "",
        "should_apply": "",
    }

    if not answer:
        return data

    percent_match = re.search(r"(\d{1,3})\s*%", answer)

    if percent_match:
        percentage = int(percent_match.group(1))
        data["match_percentage"] = max(0, min(100, percentage))

    verdict_match = re.search(
        r"Verdict\s*:\s*[-\s]*(Strong Match|Good Match|Average Match|Weak Match|Not Recommended)",
        answer,
        re.IGNORECASE,
    )

    if verdict_match:
        data["verdict"] = verdict_match.group(1)

    apply_match = re.search(
        r"Should You Apply\s*:\s*[-\s]*(Yes|Maybe|No)",
        answer,
        re.IGNORECASE,
    )

    if apply_match:
        data["should_apply"] = apply_match.group(1)

    return data


def extract_section_items(answer: str, section_name: str) -> List[str]:
    """
    Extract bullet items from model output section.

    Example section:
    Matching Skills:
    - Python
    - React
    """

    if not answer:
        return []

    pattern = rf"{re.escape(section_name)}\s*:\s*(.*?)(?=\n[A-Z][A-Za-z\s\-\/]+:\s*|$)"

    match = re.search(pattern, answer, re.IGNORECASE | re.DOTALL)

    if not match:
        return []

    section_text = match.group(1).strip()

    items = []

    for line in section_text.splitlines():
        line = line.strip()

        if not line:
            continue

        line = re.sub(r"^[-•*\d\.\)\s]+", "", line).strip()

        if line:
            items.append(line)

    return items


def build_structured_resume_result(answer: str, parsed: Dict) -> Dict:
    return {
        "match_percentage": parsed.get("match_percentage", 0),
        "verdict": parsed.get("verdict", ""),
        "should_apply": parsed.get("should_apply", ""),
        "matching_skills": extract_section_items(answer, "Matching Skills"),
        "missing_skills": extract_section_items(answer, "Missing Skills"),
        "weak_areas": extract_section_items(answer, "Weak Areas"),
        "resume_improvement_suggestions": extract_section_items(
            answer,
            "Resume Improvement Suggestions",
        ),
        "best_alternative_roles": extract_section_items(
            answer,
            "Best Alternative Roles",
        ),
    }


# ---------------------------------------------------------
# Main analyzer
# ---------------------------------------------------------

def analyze_resume_against_jd_with_rag(
    base_session_id: str,
    resume_text: str,
    jd_text: str,
    resume_filename: str = "resume",
    jd_filename: str = "job_description",
) -> Dict:
    resume_text = clean_text(resume_text)
    jd_text = clean_text(jd_text)

    if not resume_text:
        return {
            "success": False,
            "answer": "Resume text is empty.",
            "data": None,
        }

    if not jd_text:
        return {
            "success": False,
            "answer": "Job description text is empty.",
            "data": None,
        }

    try:
        index_result = build_resume_rag_index(
            base_session_id=base_session_id,
            resume_text=resume_text,
            jd_text=jd_text,
            resume_filename=resume_filename,
            jd_filename=jd_filename,
        )

        if not index_result.get("success"):
            return {
                "success": False,
                "answer": index_result.get("message", "Resume RAG indexing failed."),
                "data": None,
                "index_result": index_result,
            }

        resume_session_id = index_result["session_id"]

        retrieve_result = retrieve_resume_rag_context(
            resume_session_id=resume_session_id,
            jd_text=jd_text,
            top_k=4,
        )

        if not retrieve_result.get("success"):
            return {
                "success": False,
                "answer": retrieve_result.get("error", "Resume RAG retrieval failed."),
                "data": None,
                "index_result": index_result,
                "retrieve_result": retrieve_result,
            }

        rag_context = retrieve_result.get("context", "")

        normal_answer = call_groq_resume_llm(
            resume_text=resume_text,
            jd_text=jd_text,
            rag_context=rag_context,
        )

        parsed = parse_match_data(normal_answer)
        structured = build_structured_resume_result(normal_answer, parsed)

        return {
            "success": True,
            "answer": normal_answer,
            "data": parsed,
            "structured": structured,
            "raw_json": "",
            "resume_text": resume_text,
            "jd_text": jd_text,
            "rag_context": rag_context,
            "resume_rag_session_id": resume_session_id,
            "retrieved_chunks": retrieve_result.get("retrieved_chunks", 0),
            "total_chunks": index_result.get("total_chunks", 0),
            "model": get_groq_model(),
            "provider": "groq",
        }

    except Exception as e:
        return {
            "success": False,
            "answer": f"Error while analyzing resume with RAG: {str(e)}",
            "data": None,
            "structured": None,
            "model": get_groq_model(),
            "provider": "groq",
        }


def analyze_resume_file_bytes_and_jd_with_rag(
    base_session_id: str,
    resume_bytes: bytes,
    resume_filename: str,
    jd_text: str = "",
    jd_bytes: Optional[bytes] = None,
    jd_filename: str = "job_description",
) -> Dict:
    """
    Main FastAPI-friendly function.

    Used by:
    POST /api/resume/gap-analysis
    """

    resume_result = extract_text_from_file_bytes(
        file_bytes=resume_bytes,
        filename=resume_filename,
    )

    if not resume_result.get("success"):
        return {
            "success": False,
            "answer": f"Resume error: {resume_result.get('error')}",
            "data": None,
            "structured": None,
            "resume_text": "",
            "jd_text": "",
        }

    final_jd_text = clean_text(jd_text)
    final_jd_filename = jd_filename or "pasted_job_description"

    if jd_bytes:
        jd_result = extract_text_from_file_bytes(
            file_bytes=jd_bytes,
            filename=jd_filename,
        )

        if not jd_result.get("success"):
            return {
                "success": False,
                "answer": f"JD error: {jd_result.get('error')}",
                "data": None,
                "structured": None,
                "resume_text": resume_result.get("text", ""),
                "jd_text": "",
            }

        final_jd_text = jd_result.get("text", "")
        final_jd_filename = jd_result.get("filename", "job_description")

    analysis = analyze_resume_against_jd_with_rag(
        base_session_id=base_session_id,
        resume_text=resume_result.get("text", ""),
        jd_text=final_jd_text,
        resume_filename=resume_result.get("filename", "resume"),
        jd_filename=final_jd_filename,
    )

    return analysis


def analyze_uploaded_resume_and_jd_with_rag(
    base_session_id: str,
    resume_file,
    jd_file=None,
    jd_text: str = "",
) -> Dict:
    """
    Old Streamlit-compatible function kept for safety.
    """

    resume_result = extract_text_from_uploaded_file(resume_file)

    if not resume_result.get("success"):
        return {
            "success": False,
            "answer": f"Resume error: {resume_result.get('error')}",
            "data": None,
            "structured": None,
            "resume_text": "",
            "jd_text": "",
        }

    final_jd_text = ""
    jd_filename = "pasted_job_description"

    if jd_file is not None:
        jd_result = extract_text_from_uploaded_file(jd_file)

        if not jd_result.get("success"):
            return {
                "success": False,
                "answer": f"JD error: {jd_result.get('error')}",
                "data": None,
                "structured": None,
                "resume_text": resume_result.get("text", ""),
                "jd_text": "",
            }

        final_jd_text = jd_result.get("text", "")
        jd_filename = jd_result.get("filename", "job_description")
    else:
        final_jd_text = jd_text

    analysis = analyze_resume_against_jd_with_rag(
        base_session_id=base_session_id,
        resume_text=resume_result.get("text", ""),
        jd_text=final_jd_text,
        resume_filename=resume_result.get("filename", "resume"),
        jd_filename=jd_filename,
    )

    return analysis